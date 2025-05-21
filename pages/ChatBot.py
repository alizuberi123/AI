import streamlit as st
import os
import re
import io
import pdfplumber
import fitz  # PyMuPDF
import docx
import pytesseract
import json
from PIL import Image
from dotenv import load_dotenv
import google.generativeai as genai
from google.generativeai.types import content_types

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from utils.rag_manager import get_rag_path, save_chat_history, load_chat_history
from utils.retrieval import HybridRetriever
from utils.memory import ConversationMemory

# Load Gemini API Key 
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Check if RAG is Selected 
if "active_rag" not in st.session_state:
    st.warning("Please select a chat from the home page.")
    st.stop()

rag_name = st.session_state.active_rag
rag_path = get_rag_path(rag_name)
db_path = os.path.join(rag_path, "chroma_db")
file_dir = os.path.join(rag_path, "files")

# Streamlit Setup 
st.set_page_config(page_title=f"MIA — {rag_name}", layout="wide")
st.title(f"MIA — {rag_name}")

# Session State 
if "chat_history" not in st.session_state:
    st.session_state.chat_history = load_chat_history(rag_name)
if "db" not in st.session_state:
    st.session_state.embed_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    # Initialize Chroma with the new API
    st.session_state.db = Chroma(
        persist_directory=db_path, 
        embedding_function=st.session_state.embed_model,
        collection_name="document_collection"
    )
if "memory" not in st.session_state:
    st.session_state.memory = ConversationMemory(rag_name)
if "retriever" not in st.session_state:
    st.session_state.retriever = HybridRetriever(st.session_state.db, st.session_state.embed_model)
if "all_text" not in st.session_state:
    st.session_state.all_text = ""
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Text Cleaner 
def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\.([A-Za-z])', r'. \1', text)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
    text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
    text = re.sub(r'(?<=\w) (?=\w)', '', text)
    return text.strip()

# Extract Diagrams 
def extract_text_and_diagrams(file):
    text_chunks = []
    diagram_chunks = []

    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            text_chunks.append(clean_text(text))

        file.seek(0)
        pdf_doc = fitz.open(stream=file.read(), filetype="pdf")

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")

                image_dir = os.path.join(file_dir, "diagrams")
                os.makedirs(image_dir, exist_ok=True)
                image_path = os.path.join(image_dir, f"{file.name}_page{page_num}_{img_index}.png")
                img_pil.save(image_path)

                ocr_text = pytesseract.image_to_string(img_pil)
                cleaned = clean_text(ocr_text)

                if not cleaned.strip():
                    try:
                        vision_model = genai.GenerativeModel("models/gemini-1.5-pro-latest")
                        gemini_result = vision_model.generate_content(
                            [content_types.ImageData.from_pil_image(img_pil),
                             "Describe the diagram in detail as if explaining it to a reader."]
                        )
                        cleaned = gemini_result.text.strip()
                    except:
                        cleaned = "Unrecognized diagram content"

                if cleaned:
                    diagram_chunks.append({
                        "text": cleaned,
                        "source": file.name,
                        "page": page_num + 1,
                        "index": img_index,
                        "image_path": image_path
                    })

    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        text_chunks.append(clean_text(text))

    elif file.type == "text/plain":
        text = file.read().decode("utf-8")
        text_chunks.append(clean_text(text))

    return text_chunks, diagram_chunks

# Chunker 
def structured_chunking(raw_text: str, filename: str):
    base_name = filename.lower().replace(".pdf", "").replace(".docx", "").replace(".txt", "")
    name_parts = re.split(r"[_\-\s]+", base_name)
    tags = [part.capitalize() for part in name_parts if part]

    section_pattern = re.compile(r'^(\d+(\.\d+)*\.?\s+.+|[A-Z][A-Z\s]{3,})$', re.MULTILINE)
    lines = raw_text.split('\n')
    sections = []
    current_section = "General"
    current_content = []

    for line in lines:
        if section_pattern.match(line.strip()):
            if current_content:
                sections.append((current_section, '\n'.join(current_content)))
                current_content = []
            current_section = line.strip()
        else:
            current_content.append(line.strip())

    if current_content:
        sections.append((current_section, '\n'.join(current_content)))

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    documents = []

    for section_title, section_text in sections:
        chunks = splitter.split_text(section_text)
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk.strip(),
                metadata={
                    "filename": filename,
                    "section": section_title.strip(),
                    "chunk_id": i,
                    "tags": ", ".join(tags),
                    "type": "text"
                }
            ))
    return documents

# Gemini Reranking and Compression
def infer_topic_from_prompt(prompt):
    model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
    try:
        upgraded_prompt = f"""
You are a domain expert categorizer.

Given the following user question, identify the most specific technical topic or document section it relates to, based only on the wording of the question.

Question:
{prompt}

Respond with just the topic name. Do not include any explanations or extra text.
"""
        return model.generate_content(upgraded_prompt).text.strip()
    except:
        return "General"

def compress_chunks(chunks, question):
    model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
    combined = "\n\n".join([doc.page_content.strip() for doc in chunks])
    
    prompt = f"""
You are summarizing content for a technical Q&A assistant that answers user questions using uploaded documents.

Here is the user's question:
{question}

And here are the top most relevant document chunks:
{combined}

Your task:
- Combine the content into a clean, concise summary focused on answering the user's question.
- Remove duplicate sentences, repeated facts, or off-topic remarks.
- Keep essential details, technical terms, and step-by-step explanations.
- If applicable, include short quotes or bullet points.
- Keep the summary under 300 words if possible.

Respond with the cleaned and compressed version of the information only. Do not restate the question or explain what you're doing.
"""
    try:
        return model.generate_content(prompt).text.strip()
    except:
        return combined

# Upload Section
st.subheader("Upload & Index Notes")
uploaded_files = st.file_uploader("Upload PDFs, DOCX, or TXT files:",
                                   type=["pdf", "docx", "txt"],
                                   accept_multiple_files=True)

if uploaded_files:
    # Filter out already processed files
    new_files = [f for f in uploaded_files if f.name not in st.session_state.processed_files]
    
    if new_files:
        with st.spinner("Processing new files..."):
            new_docs = []
            os.makedirs(file_dir, exist_ok=True)  # Ensure RAG's /files exists

            for file in new_files:
                # Save uploaded file to RAG directory 
                save_path = os.path.join(file_dir, file.name)
                with open(save_path, "wb") as f:
                    f.write(file.read())
                file.seek(0)  # Reset file pointer for re-reading

                # Process content from file 
                text_blocks, diagram_blocks = extract_text_and_diagrams(file)

                for text in text_blocks:
                    st.session_state.all_text += text + "\n\n"
                    section_docs = structured_chunking(text, file.name)
                    for doc in section_docs:
                        doc.metadata["source"] = file.name
                        new_docs.append(doc)

                for diag in diagram_blocks:
                    new_docs.append(Document(
                        page_content=diag["text"],
                        metadata={
                            "source": diag["source"],
                            "page": diag["page"],
                            "chunk_id": diag["index"],
                            "section": "Diagram OCR",
                            "type": "diagram",
                            "image_path": diag["image_path"]
                        }
                    ))
                
                # Add file to processed set
                st.session_state.processed_files.add(file.name)

            # Add docs in batches 
            def add_in_batches(docs, batch_size=1000):
                for i in range(0, len(docs), batch_size):
                    st.session_state.db.add_documents(docs[i:i+batch_size])

            if new_docs:
                add_in_batches(new_docs)
                try:
                    st.session_state.db.persist()
                except:
                    pass  # Ignore if persist method doesn't exist
                
                # Reinitialize the retriever with updated database
                st.session_state.retriever = HybridRetriever(st.session_state.db, st.session_state.embed_model)
                
                st.success(f"Successfully processed {len(new_files)} new file(s)!")
            else:
                st.warning("No readable content found in new files.")
    else:
        st.info("All uploaded files have already been processed.")

# Memory Settings
with st.expander("🧠 Memory Settings"):
    col1, col2 = st.columns(2)
    with col1:
        enable_memory = st.toggle("Use Conversation Memory", value=True)
        memory_weight = st.slider("Memory Importance", min_value=0.0, max_value=1.0, value=0.3, step=0.1)
    with col2:
        use_hybrid_search = st.toggle("Use Hybrid Search", value=True)
        use_query_expansion = st.toggle("Use Query Expansion", value=True)
        
    # Thread Creation
    if st.button("Create New Conversation Thread"):
        thread_name = st.text_input("Name for the new thread:", value=f"{rag_name}_fork")
        if thread_name and thread_name.strip():
            new_thread = st.session_state.memory.create_thread(thread_name)
            if new_thread:
                st.success(f"Created new thread: {new_thread}")
                st.session_state.active_rag = new_thread
                st.rerun()
            else:
                st.error("Failed to create thread")

# Search History
with st.expander("🔍 Search Conversation History"):
    search_query = st.text_input("Search previous conversations:")
    if search_query:
        results = st.session_state.memory.search_history(search_query)
        if results:
            st.write(f"Found {len(results)} matching exchanges:")
            for i, msg in enumerate(results):
                st.write(f"**Q{i+1}:** {msg['question']}")
                st.write(f"**A{i+1}:** {msg['answer']}")
                st.divider()
        else:
            st.info("No matching conversations found.")

# Chat Interface 
if st.session_state.db:
    st.subheader("Ask Your Questions")

    # Display chat history
    for msg in st.session_state.chat_history:
        with st.chat_message("user"):
            st.markdown(msg["question"])
        with st.chat_message("assistant"):
            st.markdown(msg["answer"])

    user_input = st.chat_input("Ask your question...")

    if user_input:
        with st.spinner("MIA is thinking..."):
            # Display user message
            with st.chat_message("user"):
                st.markdown(user_input)
                
            # Get conversation memory
            memory_context = ""
            if enable_memory:
                memory_context = st.session_state.memory.get_contextual_summary(user_input)
                if memory_context:
                    st.caption(f"🧠 Using conversation memory")
            
            # Infer topic
            inferred_topic = infer_topic_from_prompt(user_input)
            st.caption(f"🔍 Inferred topic: `{inferred_topic}`")
            
            # Retrieve relevant documents using advanced retrieval
            if use_hybrid_search:
                if use_query_expansion:
                    docs = st.session_state.retriever.retrieval_with_expansion(user_input)
                else:
                    docs = st.session_state.retriever.get_relevant_documents(user_input)
            else:
                # Fallback to standard retrieval
                retriever = st.session_state.db.as_retriever()
                docs = retriever.get_relevant_documents(user_input)
                docs = rerank_chunks_with_gemini(user_input, docs)
                
            # Compress chunks
            context = compress_chunks(docs, user_input)
            
            # Construct the prompt with memory if available
            prompt = f"""
You are a helpful, knowledgeable assistant answering a user's question based on their uploaded technical documents and diagrams, with some general knowledge.

Your responsibilities:
- Primarily base your answer on the provided context from the documents.
- If the answer is not fully present in the context, you can supplement with your general knowledge.
- If the answer is completely absent from the context, you can provide a general answer but mention that it's not from the uploaded documents.
- Use plain language, but include technical terms when relevant.
- If you're using information from outside the documents, make it clear which parts are from the documents and which are from general knowledge.

Guidelines for detailed answers:
- Provide comprehensive explanations with specific examples when possible
- Include relevant technical details, specifications, or measurements
- Explain key concepts and terminology
- If applicable, mention best practices or industry standards
- Include step-by-step explanations for processes or procedures
- Compare and contrast different approaches or solutions when relevant
- Mention potential challenges or considerations
- Include practical applications or use cases

Format:
- Write 2-3 well-structured paragraphs for general questions
- Use bullet points for lists, steps, or key points
- Include relevant technical specifications or measurements
- Do NOT repeat the user's question in the answer
- If using information from outside the documents, add a note like "(This is general knowledge)" after that specific part
- For technical topics, include relevant terminology and explain complex concepts
"""

            # Add memory context if available
            if memory_context:
                prompt += f"""
Conversation memory:
{memory_context}
"""

            # Add document context
            prompt += f"""
Context from documents:
{context}

User Question:
{user_input}

Answer:
""".strip()

            model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
            stream = model.generate_content(prompt, stream=True)

            full_response = ""
            with st.chat_message("assistant"):
                box = st.empty()
                for chunk in stream:
                    if chunk.text:
                        full_response += chunk.text
                        box.markdown(full_response)

            # Save to both history systems
            new_msg = {
                "question": user_input,
                "answer": full_response,
                "sources": [doc.metadata for doc in docs]
            }
            st.session_state.chat_history.append(new_msg)
            save_chat_history(rag_name, st.session_state.chat_history)
            
            # Save to memory system
            st.session_state.memory.add_message(user_input, full_response, [doc.metadata for doc in docs])

        # Show source chunks
        with st.expander("📎 Sources Used"):
            for i, doc in enumerate(docs[:4]):
                meta = doc.metadata
                st.markdown(f"**Chunk {i+1}** — Section: `{meta.get('section', 'Unknown')}` | Type: `{meta.get('type', 'text')}` | File: `{meta.get('source', 'Unknown')}`")
                st.write(doc.page_content.strip())
                if meta.get("type") == "diagram" and meta.get("image_path") and os.path.exists(meta["image_path"]):
                    st.image(meta["image_path"], caption="Diagram Image")

    # Summary Section
    st.divider()
    st.subheader("Summarize All Uploaded Notes")
    if st.button("Generate Summary"):
        model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
        with st.spinner("Summarizing..."):
            summary = model.generate_content(f"Summarize the following notes:\n\n{st.session_state.all_text}").text.strip()
            st.success("✅ Summary")
            st.write(summary)
else:
    st.info("👆 Upload notes to begin chatting.")
